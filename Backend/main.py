"""
Warasat Backend - FastAPI Application

Unified backend for Islamic inheritance calculator and chatbot.
Provides endpoints for:
- Inheritance calculation
- Excel/PDF report generation  
- Chatbot question answering
- User authentication (login/signup)
- Real-time chat between users and Ulema
"""

import io
import json
from typing import List, Optional
from datetime import datetime

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from inheritance_calculator import calculate_inheritance
from chat_system import (
    get_user_from_token,
    create_user,
    authenticate_user,
    get_all_ulemas,
    start_or_get_chat,
    get_chat_details,
    get_chat_messages,
    send_message,
    get_ulema_chats,
    get_user_chats,
    connection_manager
)

# Lazy import for chatbot to allow API to work without chatbot dependencies
_chatbot = None

def get_chatbot():
    """Lazy load chatbot module"""
    global _chatbot
    if _chatbot is None:
        try:
            from chatbot import get_chatbot as _get_chatbot
            _chatbot = _get_chatbot()
        except ImportError as e:
            raise ImportError(
                f"Chatbot dependencies not installed. Install with: "
                f"pip install sentence-transformers pinecone-client openai PyMuPDF. "
                f"Error: {e}"
            )
    return _chatbot

# Initialize FastAPI app
app = FastAPI(
    title="Warasat API",
    description="Islamic Inheritance Calculator and Chatbot API",
    version="2.0.0"
)

# CORS middleware for frontend compatibility
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Configure appropriately for production
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============== Pydantic Models ==============

class Heir(BaseModel):
    """Model for individual heir"""
    relation: str
    category: Optional[List[str]] = None
    limit: Optional[int] = None
    val: int = 0
    amount: Optional[float] = None


class InheritanceRequest(BaseModel):
    """Request model for inheritance calculation"""
    heir_list: List[Heir]
    total_amount: float


class InheritanceResponse(BaseModel):
    """Response model for inheritance calculation"""
    status: str
    heir_list: List[dict]
    total_amount: float


class ReportRequest(BaseModel):
    """Request model for report generation"""
    total_amount: float
    funeral_expenses: Optional[float] = 0
    mehr: Optional[float] = 0
    debt: Optional[float] = 0
    will: Optional[float] = 0
    currency: Optional[str] = "USD"
    gender: Optional[str] = "male"
    heir_list: List[dict]


class ChatRequest(BaseModel):
    """Request model for chatbot"""
    query: str


class ChatResponse(BaseModel):
    """Response model for chatbot"""
    answer: str


# ============== Auth Models ==============

class SignupRequest(BaseModel):
    """Request model for user signup"""
    name: str
    email: str
    password: str


class LoginRequest(BaseModel):
    """Request model for login"""
    email: str
    password: str
    user_type: str = "user"


class AuthResponse(BaseModel):
    """Response model for auth operations"""
    success: bool
    message: Optional[str] = None
    user: Optional[dict] = None
    token: Optional[str] = None


# ============== Chat System Models ==============

class StartChatRequest(BaseModel):
    """Request model for starting a chat"""
    ulema_id: str
    inheritance_data: Optional[str] = None


class SendMessageRequest(BaseModel):
    """Request model for sending a message"""
    content: str


# ============== Helper Functions ==============

async def get_current_user(authorization: Optional[str] = Header(None)):
    """Dependency to get current authenticated user"""
    user = get_user_from_token(authorization)
    if not user:
        raise HTTPException(status_code=401, detail="Invalid or missing authentication token")
    return user


# ============== Authentication Endpoints ==============

@app.post("/auth/signup", response_model=AuthResponse)
async def signup(request: SignupRequest):
    """
    Register a new user account (regular users only, not Ulema)
    """
    user = create_user(
        name=request.name,
        email=request.email,
        password=request.password,
        user_type="user"  # Only allow regular users
    )
    
    if not user:
        raise HTTPException(status_code=400, detail="Email already registered or invalid data")
    
    return AuthResponse(
        success=True,
        message="User registered successfully"
    )


@app.post("/auth/login", response_model=AuthResponse)
async def login(request: LoginRequest):
    """
    Login with email and password
    Returns JWT token for authenticated requests
    """
    result = authenticate_user(
        email=request.email,
        password=request.password,
        user_type=request.user_type
    )
    
    if not result:
        raise HTTPException(status_code=401, detail="Invalid credentials")
    
    user, token = result
    
    return AuthResponse(
        success=True,
        user={
            "id": user.id,
            "name": user.name,
            "email": user.email,
            "user_type": user.user_type
        },
        token=token
    )


@app.get("/auth/verify")
async def verify_auth(user: dict = Depends(get_current_user)):
    """
    Verify if current token is valid
    """
    return {
        "success": True,
        "message": "Token is valid",
        "user": {
            "user_id": user["user_id"],
            "user_type": user["user_type"]
        }
    }


# ============== Ulema Endpoints ==============

@app.get("/ulemas")
async def list_ulemas(user: dict = Depends(get_current_user)):
    """
    Get list of all available Ulema
    """
    ulemas = get_all_ulemas()
    return {"success": True, "ulemas": ulemas}


# ============== Chat Endpoints ==============

@app.post("/chats/start")
async def start_chat(request: StartChatRequest, user: dict = Depends(get_current_user)):
    """
    Start a new chat with an Ulema or get existing chat ID
    """
    chat_id = start_or_get_chat(
        user_id=user["user_id"],
        ulema_id=request.ulema_id,
        inheritance_data=request.inheritance_data
    )
    return {"success": True, "chat_id": chat_id}


@app.get("/chats/{chat_id}")
async def get_chat(chat_id: str, user: dict = Depends(get_current_user)):
    """
    Get details of a specific chat
    """
    chat = get_chat_details(chat_id, user["user_id"])
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found or unauthorized")
    
    return {"success": True, "chat": chat}


@app.get("/chats/{chat_id}/messages")
async def get_messages(chat_id: str, user: dict = Depends(get_current_user)):
    """
    Get all messages in a chat
    """
    messages = get_chat_messages(chat_id, user["user_id"], user["user_type"])
    if messages is None:
        raise HTTPException(status_code=404, detail="Chat not found or unauthorized")
    
    return {"success": True, "messages": messages}


@app.post("/chats/{chat_id}/messages")
async def post_message(chat_id: str, request: SendMessageRequest, user: dict = Depends(get_current_user)):
    """
    Send a message in a chat
    """
    message = send_message(
        chat_id=chat_id,
        sender_id=user["user_id"],
        sender_type=user["user_type"],
        content=request.content
    )
    
    if not message:
        raise HTTPException(status_code=404, detail="Chat not found or unauthorized")
    
    # Broadcast to WebSocket connections
    await connection_manager.broadcast_to_chat(chat_id, {
        "type": "new_message",
        "message": message
    })
    
    return {"success": True, "message": message}


@app.get("/ulema/chats")
async def get_ulema_chat_list(user: dict = Depends(get_current_user)):
    """
    Get all chats for the authenticated Ulema
    """
    if user["user_type"] != "ulema":
        raise HTTPException(status_code=403, detail="Only Ulema can access this endpoint")
    
    chats = get_ulema_chats(user["user_id"])
    return {"success": True, "chats": chats}


@app.get("/user/chats")
async def get_user_chat_list(user: dict = Depends(get_current_user)):
    """
    Get all chats for the authenticated user
    """
    chats = get_user_chats(user["user_id"])
    return {"success": True, "chats": chats}


# ============== WebSocket Endpoint ==============

@app.websocket("/ws/chat/{chat_id}")
async def websocket_chat(websocket: WebSocket, chat_id: str):
    """
    WebSocket endpoint for real-time chat
    """
    # Get token from query params
    token = websocket.query_params.get("token")
    if not token:
        await websocket.close(code=4001)
        return
    
    user = get_user_from_token(f"Bearer {token}")
    if not user:
        await websocket.close(code=4001)
        return
    
    # Verify user has access to this chat
    chat = get_chat_details(chat_id, user["user_id"])
    if not chat:
        await websocket.close(code=4003)
        return
    
    # Connect
    await connection_manager.connect(websocket, chat_id)
    
    try:
        while True:
            # Receive message
            data = await websocket.receive_json()
            
            if data.get("type") == "message":
                # Save message to database
                message = send_message(
                    chat_id=chat_id,
                    sender_id=user["user_id"],
                    sender_type=user["user_type"],
                    content=data.get("content", "")
                )
                
                if message:
                    # Broadcast to all connections in this chat
                    await connection_manager.broadcast_to_chat(chat_id, {
                        "type": "new_message",
                        "message": message
                    })
    except WebSocketDisconnect:
        connection_manager.disconnect(websocket, chat_id)


# ============== Inheritance Calculator Endpoints ==============

@app.post("/inheritance-calculator-2", response_model=InheritanceResponse)
async def calculate_inheritance_shares(request: InheritanceRequest):
    """
    Calculate inheritance shares for given heirs
    
    Takes a list of heirs with their counts and total wealth,
    returns each heir's calculated share amount.
    """
    try:
        # Convert Pydantic models to dicts
        heir_list = [heir.model_dump() for heir in request.heir_list]
        total_amount = request.total_amount
        
        # Calculate inheritance
        result_heir_list = calculate_inheritance(total_amount, heir_list)
        
        return InheritanceResponse(
            status="success",
            heir_list=result_heir_list,
            total_amount=total_amount
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== Report Generation Endpoints ==============

@app.post("/inheritance-calculation-xlsx")
async def generate_excel_report(request: ReportRequest):
    """
    Generate Excel report for inheritance calculation
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, Border, Side
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Inheritance Calculation"
        
        # Styles
        header_font = Font(bold=True, size=14)
        title_font = Font(bold=True, size=16)
        center_align = Alignment(horizontal="center", vertical="center")
        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin")
        )
        
        # Title
        ws.merge_cells("A1:G1")
        ws["A1"] = "Islamic Inheritance Calculation Report"
        ws["A1"].font = title_font
        ws["A1"].alignment = center_align
        
        # Date
        ws["A2"] = f"Date: {datetime.now().strftime('%d-%m-%Y')}"
        
        # Asset Details Section
        ws["A4"] = "Asset Details"
        ws["A4"].font = header_font
        
        row = 5
        asset_details = {
            "Total Amount": request.total_amount,
            "Funeral & Burial Expenses": request.funeral_expenses,
            "Haq Mehr": request.mehr,
            "Debt & Loans": request.debt,
            "Will": request.will,
        }
        
        distributable = (request.total_amount - (request.funeral_expenses or 0) - 
                        (request.mehr or 0) - (request.debt or 0) - (request.will or 0))
        asset_details["Distributable Amount"] = distributable
        
        for key, value in asset_details.items():
            if value and value > 0:
                ws[f"A{row}"] = key
                ws[f"B{row}"] = f"{request.currency} {value:,.2f}"
                ws[f"A{row}"].border = thin_border
                ws[f"B{row}"].border = thin_border
                row += 1
        
        # Heir Details Section
        row += 2
        ws[f"A{row}"] = "Heir Details"
        ws[f"A{row}"].font = header_font
        row += 1
        
        # Headers
        headers = ["Count", "Relation", "Category", "Share Amount"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = center_align
            cell.border = thin_border
        row += 1
        
        # Heir data
        for heir in request.heir_list:
            ws.cell(row=row, column=1, value=heir.get("val", 0)).border = thin_border
            ws.cell(row=row, column=2, value=heir.get("relation", "")).border = thin_border
            category = heir.get("category", [])
            ws.cell(row=row, column=3, value=category[-1] if category else "").border = thin_border
            amount = heir.get("amount", 0)
            ws.cell(row=row, column=4, value=f"{request.currency} {amount:,.2f}").border = thin_border
            row += 1
        
        # Adjust column widths
        ws.column_dimensions["A"].width = 15
        ws.column_dimensions["B"].width = 25
        ws.column_dimensions["C"].width = 15
        ws.column_dimensions["D"].width = 20
        
        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=Inheritance-Calculation.xlsx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/inheritance-calculation-pdf")
async def generate_pdf_report(request: ReportRequest):
    """
    Generate PDF report for inheritance calculation
    Uses python-docx to create a Word document, then convert or serve directly
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading("Islamic Inheritance Calculation Report", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # Asset Details
        doc.add_heading("Asset Details", level=1)
        
        distributable = (request.total_amount - (request.funeral_expenses or 0) - 
                        (request.mehr or 0) - (request.debt or 0) - (request.will or 0))
        
        asset_table = doc.add_table(rows=1, cols=2)
        asset_table.style = "Table Grid"
        
        asset_details = [
            ("Total Amount", request.total_amount),
            ("Funeral & Burial Expenses", request.funeral_expenses),
            ("Haq Mehr", request.mehr),
            ("Debt & Loans", request.debt),
            ("Will", request.will),
            ("Distributable Amount", distributable),
        ]
        
        for label, value in asset_details:
            if value and value > 0:
                row = asset_table.add_row()
                row.cells[0].text = label
                row.cells[1].text = f"{request.currency} {value:,.2f}"
        
        # Remove first empty row
        asset_table._tbl.remove(asset_table.rows[0]._tr)
        
        doc.add_paragraph()
        
        # Heir Details
        doc.add_heading("Heir Details", level=1)
        
        heir_table = doc.add_table(rows=1, cols=4)
        heir_table.style = "Table Grid"
        
        # Headers
        headers = ["Count", "Relation", "Category", "Share Amount"]
        header_cells = heir_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Data rows
        for heir in request.heir_list:
            row = heir_table.add_row()
            row.cells[0].text = str(heir.get("val", 0))
            row.cells[1].text = heir.get("relation", "")
            category = heir.get("category", [])
            row.cells[2].text = category[-1] if category else ""
            amount = heir.get("amount", 0)
            row.cells[3].text = f"{request.currency} {amount:,.2f}"
        
        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=Inheritance-Calculation.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== Chatbot Endpoints ==============

@app.post("/get_answer", response_model=ChatResponse)
async def get_chatbot_answer(request: ChatRequest):
    """
    Get answer from the inheritance chatbot
    
    Uses RAG (Retrieval Augmented Generation) to answer
    questions about Islamic inheritance law.
    """
    try:
        if not request.query:
            raise HTTPException(status_code=400, detail="Query is required")
        
        chatbot = get_chatbot()
        answer = chatbot.get_answer(request.query)
        
        return ChatResponse(answer=answer)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== Health Check ==============

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "service": "warasat-api"}


@app.get("/")
async def root():
    """Root endpoint with API info"""
    return {
        "name": "Warasat API",
        "version": "2.0.0",
        "description": "Islamic Inheritance Calculator, Chatbot, and Chat System",
        "endpoints": {
            "inheritance": "/inheritance-calculator-2",
            "excel_report": "/inheritance-calculation-xlsx",
            "pdf_report": "/inheritance-calculation-pdf",
            "chatbot": "/get_answer",
            "health": "/health",
            "auth": {
                "signup": "/auth/signup",
                "login": "/auth/login",
                "verify": "/auth/verify"
            },
            "chat": {
                "ulemas": "/ulemas",
                "start_chat": "/chats/start",
                "get_chat": "/chats/{chat_id}",
                "messages": "/chats/{chat_id}/messages",
                "ulema_chats": "/ulema/chats",
                "user_chats": "/user/chats",
                "websocket": "/ws/chat/{chat_id}"
            }
        }
    }


# ============== Main Entry Point ==============

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8080)

